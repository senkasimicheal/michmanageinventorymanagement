from flask import Flask, render_template, url_for, send_from_directory, request, flash, redirect, session, make_response, jsonify, current_app
import secrets
from flask_mail import Message
import threading
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from docx import Document
from pymongo import MongoClient, ASCENDING, DESCENDING
import secrets
import bcrypt
from datetime import datetime, timedelta, timezone
import calendar
import pytz
import pandas as pd 
from io import BytesIO
import json
from bson.objectid import ObjectId
import cv2
import numpy as np
import io
import base64
import random
import os
from werkzeug.utils import secure_filename
from gridfs import GridFS
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Image
from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from zipfile import ZipFile
import tempfile
import string
import qrcode
import time
from docx2pdf import convert
import PyPDF2
import gc
from collections import defaultdict
import barcode
from PIL import Image, ImageDraw, ImageFont
from barcode.writer import ImageWriter

# Import blueprints
from user_register import registration
from user_login import login
from documentation import app_documentation
from stock_management import stockManagement
from manager_account_setup import managerAccountSetup
from other_user_accounts_mngt import otherUserAccounts
from user_rights import userRights
from logs import logs
from admin import Admin
from accounting import accounting
from utils import get_mongo_client, get_db_and_fs, get_mail_instance, send_async_email

app = Flask(__name__, static_folder='static')
app.secret_key = secrets.token_hex(16)
scheduler = BackgroundScheduler()

app.config.update(
    MAIL_SERVER='smtp.sendgrid.net',
    MAIL_PORT=587,
    MAIL_USERNAME='apikey',
    MAIL_PASSWORD='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',
    MAIL_USE_TLS=True,
    MAIL_USE_SSL=False
)

# Initialize Mail
mail = get_mail_instance(app)

utc = pytz.UTC

# Register blueprints
app.register_blueprint(registration)
app.register_blueprint(login)
app.register_blueprint(app_documentation)
app.register_blueprint(stockManagement)
app.register_blueprint(managerAccountSetup)
app.register_blueprint(otherUserAccounts)
app.register_blueprint(userRights)
app.register_blueprint(logs)
app.register_blueprint(Admin)
app.register_blueprint(accounting)

@app.after_request
def after_request(response):
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate, max-age=0"
    response.headers["Expires"] = '0'
    response.headers["Pragma"] = "no-cache"
    return response

@app.route("/")
def index():
    return render_template('index.html')

@app.route("/rate_us")
def rate_us():
    return redirect('https://www.google.com/search?client=ms-android-transsion&sca_esv=033c270f53017be6&sxsrf=ADLYWILRhWsAAhDBzis_Uv0L4b0A9HFehA:1724743888012&q=mich+manage+kampala+reviews&uds=ADvngMgfPPaym3hQlwI19QyQZ3bfoUA6G3RfDiV_RezhubPzsBvD58rj02PtNO6qYPH6K_9RFI5-ZhN5v91U54cbknXUtJzjUgqxKE-3cpa9Vms2l6e0bv3ppik1dC3mvdYVH3ik3CpYoSoB0hEMAtxL3Cwn1EAx15RwHCoISeVaQpGnU3k1fwAIj-RTls2NpzCeBju4vsZQoiielCBsJpLKu_rMRxQG-n4wHnIFsEbC83jGq-Ub3MGyp0icXviM9j7DV4lLKjK7ADwR8xVRjdwRIfJFKYbFa-KVcgsBHngLB6lmE3FpIybfDEwqp2U5NJ7ITxSm3nf5HMlZIX6a2yGUZgOkVsS377_6GL6B2NWkv0rXPonn--9yHyYP7rS_yhAfV7D10slP0OctugY-Ceqo8A3cS5pV2BKeZF-tTrGiZtgtk9rpCitKqVs6mmYhBrQAae0G7Zjn&si=ACC90nwjPmqJHrCEt6ewASzksVFQDX8zco_7MgBaIawvaF4-7uH_XqoShGoDxiuBCUkUzWczFmimLaxlPIvUzNC0JYqo2RjwLk-dmlvITGroH7G62nlGKf4wXc5GWWRkcgQutIDlcU5lIbDDTqF6wP9HoxTvS79K-w%3D%3D&sa=X&ved=2ahUKEwj91pfW05SIAxWPBdsEHcSrI-gQk8gLegQIIBAB&ictx=1&biw=360&bih=680&dpr=2')

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/', code=303)

@app.route('/logout-admin')
def logout_admin():
    session.clear()
    return redirect('/admin', code=303)

@app.route('/googlee9cdc37dc478e7a2.html')
def google_verification():
    return render_template('googlee9cdc37dc478e7a2.html')

@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory(app.static_folder, request.path[1:])

def add_password_to_pdf(pdf_path, password):
    output_path = pdf_path.replace('.pdf', '_protected.pdf')
    pdf_writer = PyPDF2.PdfWriter()
    pdf_reader = PyPDF2.PdfReader(pdf_path)
    
    for page_num in range(len(pdf_reader.pages)):
        pdf_writer.add_page(pdf_reader.pages[page_num])
    
    pdf_writer.encrypt(user_pwd=password, owner_pwd=None, use_128bit=True)
    
    with open(output_path, 'wb') as f:
        pdf_writer.write(f)
    
    return output_path

def convert_docx_to_pdf(docx_path):
    convert(docx_path)
    pdf_path = docx_path.replace('.docx', '.pdf')
    return pdf_path

##########SEND PAYMENT REMINDERS###########
def send_payment_financial_reminders():
    current_day_of_week = datetime.now().weekday()
    if current_day_of_week != 3 or current_day_of_week != 4:
        return
    db, fs = get_db_and_fs()
    send_emails = db.send_emails.find_one({'emails': "yes"},{'emails': 1})
    accounts = list(db.transaction_finance_accounts.find())
    for account in accounts:
        if account['amount_demanded'] > 0:
            email = account.get('email')
            if email and email.strip():
                user_email = account['email']
                #Sending reminder message
                if send_emails is not None:
                    msg = Message('Payment Reminders - Mich Manage', 
                    sender='michmanage@outlook.com', 
                    recipients=[user_email])
                    msg.html = f"""
                    <html>
                    <body>
                    <p>Dear {account['client_name']},</p>
                    <p>This is a friendly reminder that a payment of <b>{ account['amount_demanded'] }</b> is currently due from <b>{ account['company_name'] }</b>.</p>
                    <p>We kindly request that you ensure your payment is processed at your earliest convenience.</p>
                    <p>If you have any questions or require further assistance, please do not hesitate to contact us.</p>
                    <p>Thank you for your prompt attention to this matter.</p>
                    <p>Best Regards,</p>
                    <p>Mich Manage</p>
                    </body>
                    </html>
                    """
                    # Send the email
                    with app.app_context():
                        thread = threading.Thread(target=send_async_email, args=[current_app._get_current_object(), msg])
                        thread.start()

##########SEND MONTHLY REPORTS###########
def send_inventory_reports():
    if datetime.now().day != 1:
        return

    db, fs = get_db_and_fs()
    send_emails = db.send_emails.find_one({'emails': "yes"},{'emails': 1})
    current_year = datetime.now().year
    current_month = datetime.now().month
    now = datetime.now()

    if current_month == 1:
        previous_month = 12
        previous_month_year = current_year - 1
    else:
        previous_month = current_month - 1
        previous_month_year = current_year

    now_without_seconds = now.replace(second=0, microsecond=0)
    first_day_previous_month = (now.replace(day=1) - timedelta(days=1)).replace(day=1)
    first_day_previous_month_str = first_day_previous_month.strftime('%B %d, %Y')
    manager_emails = db.managers.find({'account_type': 'Enterprise Resource Planning'}, {'manager_email': 1, 'name': 1})
    
    for manager_email in manager_emails:
        email = manager_email['manager_email']
        # Query the registered collection for the username associated with the manager's email
        registered_doc = db.registered_managers.find_one({'email': email, 'company_name': manager_email['name']})
        if registered_doc:
            company_name = registered_doc['company_name']

            previous_month_paid = datetime.now().month - 1 if datetime.now().month > 1 else 12

            pipeline = [
                {
                    '$match': {
                        'company_name': company_name,
                        'stockDate': {
                            '$gte': datetime(current_year, previous_month_paid, 1).replace(tzinfo=utc),
                            '$lt': datetime(current_year, previous_month_paid + 1, 1).replace(tzinfo=utc) if previous_month_paid < 12 else datetime(current_year + 1, 1, 1).replace(tzinfo=utc)
                        }
                    }
                },
                {
                    '$group': {
                        '_id': {'itemName': '$itemName','stockDate': '$stockDate'},
                        'totalRevenue': {'$sum': '$revenue'},
                        'quantitysold': {'$sum': '$quantity'}
                    }
                },
                {
                    '$lookup': {
                        'from': 'inventories',
                        'let': {'itemName': '$_id.itemName', 'stockDate': '$_id.stockDate'},
                        'pipeline': [
                            {
                                '$match': {
                                    '$expr': {
                                        '$and': [
                                            {'$eq': ['$itemName', '$$itemName']},
                                            {'$eq': ['$company_name', company_name]},
                                            { '$gte': ['$stockDate', datetime(current_year, previous_month_paid, 1).replace(tzinfo=utc)] },
                                            { '$lt': ['$stockDate', datetime(current_year, previous_month_paid + 1, 1).replace(tzinfo=utc) if previous_month_paid < 12 else datetime(current_year + 1, 1, 1).replace(tzinfo=utc)] }
                                        ]
                                    }
                                }
                            },
                            {
                                '$project': {
                                    '_id': 0,
                                    'quantity': 1,
                                    'unitPrice': 1,
                                    'stockDate': 1,
                                    'totalPrice': 1
                                }
                            }
                        ],
                        'as': 'inventoryDetails'
                    }
                }
            ]

            revenue_info = list(db.stock_sales.aggregate(pipeline))

            item_names = []
            quantities_sold = []
            quantities_stocked = []
            total_revenues = []
            total_prices = []
            profits = []

            for record in revenue_info:
                item_names.append(record['_id']['itemName'])
                quantities_sold.append(record['quantitysold'])
                total_revenues.append(record['totalRevenue'])

                # Initialize variables for each iteration
                quantities_stocked_iter = 0
                total_price_iter = 0
                profit_iter = 0

                # Check if 'inventoryDetails' is in record and has the necessary structure
                if 'inventoryDetails' in record and record['inventoryDetails']:
                    quantity_stocked = record['inventoryDetails'][0].get('quantity', 0)
                    total_price_iter = record['inventoryDetails'][0].get('totalPrice', 0)
                    profit_iter = record['totalRevenue'] - total_price_iter

                # Append values for this iteration to the lists
                quantities_stocked.append(quantities_stocked_iter)
                total_prices.append(total_price_iter)
                profits.append(profit_iter)

            # Create the DataFrame
            df_ungrouped = pd.DataFrame({
                'Item Name': item_names,
                'Quantity Sold': quantities_sold,
                'Quantity Stocked': quantities_stocked,
                'Total Revenue': total_revenues,
                'Total Price': total_prices,
                'Profit': profits
            })

            # Group by 'Item Name' and aggregate using sum
            df = df_ungrouped.groupby('Item Name', as_index=False).agg({
                'Quantity Sold': 'sum',
                'Quantity Stocked': 'sum',
                'Total Revenue': 'sum',
                'Total Price': 'sum',
                'Profit': 'sum'
            })

            profitableItems = []
            profits = []
            top_profitable_items = df[df['Profit'] > 0].sort_values(by='Profit', ascending=False).head(10)

            if not top_profitable_items.empty:
                for index, row in top_profitable_items.iterrows():
                    profitableItems.append(row['Item Name'])
                    profits.append(row['Profit'])
            top10profits = list(zip(profitableItems, profits))

            unprofitableItems = []
            losses = []
            top_unprofitable_items = df[df['Profit'] < 0].sort_values(by='Profit', ascending=False).head(10)

            if not top_unprofitable_items.empty:
                for index, row in top_unprofitable_items.iterrows():
                    unprofitableItems.append(row['Item Name'])
                    losses.append(row['Profit'])
            top10losses = list(zip(unprofitableItems, losses))

            ##total revenue
            revenueItems = []
            revenue = []
            if not df.empty:
                top_revenue = df.sort_values(by='Total Revenue', ascending=False).head(10)
                for index, row in top_revenue.iterrows():
                    revenueItems.append(row['Item Name'])
                    revenue.append(row['Total Revenue'])
            top10revenues = list(zip(revenueItems, revenue))

            ##Quantity sold
            soldItems = []
            soldQuantity = []
            if not df.empty:
                quantity_sold = df.sort_values(by='Quantity Sold', ascending=False).head(10)
                for index, row in quantity_sold.iterrows():
                    soldItems.append(row['Item Name'])
                    soldQuantity.append(row['Quantity Sold'])
            top10SoldItems = list(zip(soldItems, soldQuantity))

            ###PROFIT TRENDS
            twelve_months_ago = datetime.now() - timedelta(days=365)
            pipeline_profits = [
                {
                    '$match': {
                        'company_name': company_name,
                        'stockDate': {'$gte': twelve_months_ago}
                    }
                },
                {
                    '$group': {
                        '_id': {'itemName': '$itemName', 'stockDate': '$stockDate'},
                        'totalRevenue': {'$sum': '$revenue'},
                        'quantitysold': {'$sum': '$quantity'}
                    }
                },
                {
                    '$lookup': {
                        'from': 'inventories',
                        'let': {'itemName': '$_id.itemName', 'stockDate': '$_id.stockDate'},
                        'pipeline': [
                            {
                                '$match': {
                                    '$expr': {
                                        '$and': [
                                            {'$eq': ['$itemName', '$$itemName']},
                                            {'$eq': ['$company_name', company_name]},
                                            {'$gte': ['$stockDate', twelve_months_ago]}
                                        ]
                                    }
                                }
                            },
                            {
                                '$project': {
                                    '_id': 0,
                                    'quantity': 1,
                                    'unitPrice': 1,
                                    'stockDate': 1,
                                    'totalPrice': 1
                                }
                            }
                        ],
                        'as': 'inventoryDetails'
                    }
                }
            ]
            profit_info = list(db.stock_sales.aggregate(pipeline_profits))

            profit_item_names = []
            profit_data = []
            profit_stock_dates = []


            for profit_record in profit_info:
                if 'inventoryDetails' in profit_record and profit_record['inventoryDetails']:
                    profit_item_names.append(profit_record['_id']['itemName'])
                    profit_data.append(profit_record['totalRevenue'] - profit_record['inventoryDetails'][0]['totalPrice'])
                    profit_stock_dates.append(profit_record['_id']['stockDate'])

            # Create the DataFrame
            profit_info_df = pd.DataFrame({
                'Item Name': profit_item_names,
                'Profit': profit_data,
                'Stock Date': profit_stock_dates
            })

            # Convert 'Stock Date' to datetime
            profit_info_df['Stock Date'] = pd.to_datetime(profit_info_df['Stock Date'])

            # Group by month and calculate the sum of profits
            monthly_profits = profit_info_df.groupby(profit_info_df['Stock Date'].dt.to_period('M'))['Profit'].sum()

            # Create a DataFrame with month names
            monthly_profits_df = pd.DataFrame({
                'Month': monthly_profits.index.to_timestamp().strftime('%B'),
                'Monthly Profit': monthly_profits
            })
            months = []
            monthlyProfits = []
            if not monthly_profits_df.empty:
                for index, row in monthly_profits_df.iterrows():
                    months.append(row['Month'])
                    monthlyProfits.append(row['Monthly Profit'])
            _12monthprofits = list(zip(months, monthlyProfits))

            # Create a new Word document
            doc = Document()
            doc.add_heading(f'Inventory Performance Report for {company_name}', 0)

            # Add the data to the document
            doc.add_paragraph(f'Date: {now_without_seconds}')
            doc.add_heading('Executive Summary', level=2)
            doc.add_paragraph(f'This report provides an overview of the inventory management activities for the period from {first_day_previous_month_str} to {calendar.month_name[previous_month]} {calendar.monthrange(previous_month_year, previous_month)[1]}, {previous_month_year}. It includes key performance indicators, and financial summaries.')
            
            doc.add_heading('Top 10 Profitable Items', level=2)
            count = 1
            for item, profit in top10profits:
                doc.add_paragraph(f'{count}. {item}: UGX {profit}')
                count +=1

            doc.add_heading('Top 10 Unprofitable Items', level=2)
            count = 1
            for item, loss in top10losses:
                doc.add_paragraph(f'{count}. {item}: UGX {loss}')
                count +=1

            doc.add_heading('Top 10 Revenue-Generating Items', level=2)
            count = 1
            for item, revenue in top10revenues:
                doc.add_paragraph(f'{count}. {item}: UGX {revenue}')
                count +=1
            
            doc.add_heading('Top 10 Most Sold Items', level=2)
            count = 1
            for item, quantity in top10SoldItems:
                doc.add_paragraph(f'{count}. {item}: {quantity}')
                count +=1

            doc.add_heading('Monthly Profit Trends', level=2)
            count = 1
            for month, monthly_profit in _12monthprofits:
                doc.add_paragraph(f'{count}. {month}: {monthly_profit}')
                count +=1

            # Save the document
            report_filename = f'{company_name}_inventory_report.docx'
            doc.save(report_filename)

            pdf_filename = convert_docx_to_pdf(report_filename)

            # Create a new Flask-Mail Message
            if send_emails is not None:
                msg = Message(
                    'Mich Manage - Monthly Inventory Performance Report',
                    sender='michmanage@outlook.com',
                    recipients=[email]
                )

                # Attach the report
                with app.open_resource(pdf_filename) as fp:
                    msg.attach(pdf_filename, "application/pdf", fp.read())

                # Set the HTML body of the email
                msg.html = f"""
                <html>
                <body>
                <p>Dear {company_name},</p>
                <p>Please find attached your monthly report.</p>
                <p>Best Regards,</p>
                <p>Mich Manage</p>
                </body>
                </html>
                """

                # Send the email
                with app.app_context():
                    thread = threading.Thread(target=send_async_email, args=[current_app._get_current_object(), msg])
                    thread.start()
                # Delete the report
                os.remove(report_filename)
                os.remove(pdf_filename)
                del df
                gc.collect()


scheduler.add_job(
    func=send_inventory_reports,
    trigger=CronTrigger(day=1, hour=9, minute=0),
    id='send_inventory_reports_job',
    name='Send reports on the 1st of every month',
    replace_existing=True
)

scheduler.add_job(
    send_payment_financial_reminders,
    CronTrigger(hour=12, minute=0),
    id='send_payment_financial_reminders_job',
    name='Run job every day at 12 PM',
    replace_existing=True
)

scheduler.start()